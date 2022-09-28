import docx
from docx.enum.text import WD_BREAK
from docx.shared import Pt


class WordFileManager:
    p_fs = 14
    text_flag = 'литература'

    def __init__(self, file_name):
        self.doc = docx.Document(file_name)
        self.is_text_changing = True

    def is_default_text(self, p):
        """Базовый ли это текст"""
        for run in p.runs:
            if (run.font.size == Pt(self.p_fs)) and (not run.font.bold):
                return True
        return False

    def exists_flag(self, text):
        """Есть ли необходимое флаг-слово в тексте"""
        return self.text_flag in text

    def set_default(self):
        """Приведение к стартовому виду всех переменных"""
        self.is_text_changing = True

    def _replace_text(self, p, old, new):
        """Замена текста"""
        if not self.is_default_text(p):
            self.is_text_changing = not self.exists_flag(p.text)
        if self.is_text_changing:
            inline = p.runs
            for i in range(len(inline)):
                inline[i].text = inline[i].text.replace(old, new)
                if inline[i].text == '\n':
                    inline[i].add_break(WD_BREAK.PAGE)

    def replace_text_paragraphs(self, old, new):
        """Замена текста в параграфах"""
        for p in self.doc.paragraphs:
            # print(p.text)
            self._replace_text(p, old, new)
        self.set_default()

    def replace_text_tables(self, old, new):
        """Замена текста в таблицах"""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_text(p, old, new)
        self.set_default()

    def replace_text(self, old, new):
        """Замена всех видов текста"""
        self.replace_text_paragraphs(old, new)
        self.replace_text_tables(old, new)
        # self.doc.add_picture()

    def print(self):
        for p in self.doc.paragraphs:
            print(p.text)

    def save(self, path):
        self.doc.save(path)
